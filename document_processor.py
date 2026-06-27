"""
Document ingestion pipeline: PDF, DOCX, TXT → chunks → index.
"""
from __future__ import annotations
import os, re, hashlib
from pathlib import Path
from typing import List, Dict, Optional, Iterator
from dataclasses import dataclass
import logging

logger = logging.getLogger(__name__)


@dataclass
class Document:
    content: str
    source: str
    doc_type: str
    metadata: Dict
    doc_id: str = ""

    def __post_init__(self):
        if not self.doc_id:
            self.doc_id = hashlib.md5(
                f"{self.source}{self.content[:100]}".encode()
            ).hexdigest()[:12]


class DocumentProcessor:
    """
    Loads and pre-processes documents from multiple formats.
    Supported: PDF, DOCX, TXT, MD, HTML
    """

    SUPPORTED = {".pdf", ".docx", ".txt", ".md", ".html", ".htm"}

    def __init__(self, min_chars: int = 50, deduplicate: bool = True):
        self.min_chars = min_chars
        self.deduplicate = deduplicate
        self._seen: set = set()

    def load_file(self, path: str) -> Optional[Document]:
        path = Path(path)
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED:
            logger.warning(f"Unsupported format: {suffix}")
            return None

        try:
            if suffix == ".pdf":
                content = self._load_pdf(path)
            elif suffix == ".docx":
                content = self._load_docx(path)
            elif suffix in (".html", ".htm"):
                content = self._load_html(path)
            else:
                content = path.read_text(encoding="utf-8", errors="replace")

            content = self._clean(content)
            if len(content) < self.min_chars:
                return None

            if self.deduplicate:
                h = hashlib.md5(content.encode()).hexdigest()
                if h in self._seen:
                    logger.debug(f"Duplicate skipped: {path.name}")
                    return None
                self._seen.add(h)

            return Document(
                content=content,
                source=str(path),
                doc_type=suffix.lstrip("."),
                metadata={"filename": path.name, "size_chars": len(content)},
            )

        except Exception as e:
            logger.error(f"Failed to load {path}: {e}")
            return None

    def load_directory(self, directory: str, recursive: bool = True) -> List[Document]:
        docs = []
        pattern = "**/*" if recursive else "*"
        for path in Path(directory).glob(pattern):
            if path.is_file() and path.suffix.lower() in self.SUPPORTED:
                doc = self.load_file(str(path))
                if doc:
                    docs.append(doc)
        logger.info(f"Loaded {len(docs)} documents from {directory}")
        return docs

    def load_text(self, text: str, source: str = "manual") -> Document:
        content = self._clean(text)
        return Document(content=content, source=source, doc_type="text", metadata={})

    # ── Loaders ───────────────────────────────────────────────────────────────

    def _load_pdf(self, path: Path) -> str:
        try:
            import pdfplumber
            with pdfplumber.open(path) as pdf:
                return "\n".join(
                    page.extract_text() or "" for page in pdf.pages
                )
        except ImportError:
            pass
        try:
            import pypdf
            reader = pypdf.PdfReader(str(path))
            return "\n".join(page.extract_text() or "" for page in reader.pages)
        except ImportError:
            raise ImportError("pip install pdfplumber or pypdf")

    def _load_docx(self, path: Path) -> str:
        try:
            from docx import Document as DocxDoc
            doc = DocxDoc(str(path))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except ImportError:
            raise ImportError("pip install python-docx")

    def _load_html(self, path: Path) -> str:
        try:
            from bs4 import BeautifulSoup
            html = path.read_text(encoding="utf-8", errors="replace")
            soup = BeautifulSoup(html, "html.parser")
            for tag in soup(["script", "style", "nav", "footer"]):
                tag.decompose()
            return soup.get_text(separator="\n")
        except ImportError:
            raise ImportError("pip install beautifulsoup4")

    @staticmethod
    def _clean(text: str) -> str:
        text = re.sub(r"\n{3,}", "\n\n", text)       # collapse blank lines
        text = re.sub(r"[ \t]{2,}", " ", text)          # collapse spaces
        text = re.sub(r"[^\S\n]+$", "", text, flags=re.M)  # trailing spaces
        return text.strip()
